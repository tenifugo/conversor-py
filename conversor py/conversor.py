import tkinter as tk
from tkinter import messagebox, ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def process_text():
    input_text = text_area.get("1.0", tk.END).strip()
    form_data = parse_form_data(input_text)
    
    if form_data:
        save_to_word(form_data)
        messagebox.showinfo("Sucesso", "Dados exportados para Word com sucesso!")
    else:
        messagebox.showerror("Erro", "Não foi possível processar os dados.")

def parse_form_data(data):
    entries = re.findall(r'(\w+)\s*\n\s*([^\n]+)', data)
    return {key: value for key, value in entries}

def format_key(key):
    question_mapping = {
        'razao_social': 'Razão Social da Empresa:',
        'cnpj': 'Número de Identificação Fiscal (CNPJ):',
        'nome_fantasia': 'Nome Fantasia:',
        'data_abertura': 'Data de Constituição/Abertura da Empresa:',
        'atividade_principal': 'Atividade Principal da Empresa:',
        'endereco': 'Endereço Completo:',
        'cep': 'CEP:',
        'cidade': 'Cidade:',
        'estado': 'Estado:',
        'site': 'Site da Empresa:',
        'arranjo_selecionado': 'O Cliente é membro de um ou mais Arranjo(s) Societário(s)?',
        'descricao_arranjo': 'Descrição do Arranjo:',
        'nome_socio1': 'Nome Completo dos Sócios/Diretores/Administradores/Representantes Legais 1:',
        'cpf_socio1': 'CPF/CNPJ 1:',
        'data_nascimento1': 'Data de Nascimento 1:',
        'percentual_societario1': 'Percentual Societário 1:',
        'us_person1': 'É US Person 1?',
        'compliance_selecionado': 'A empresa possui departamento de Compliance?',
        'compliance_officer': 'Nome do(a) Compliance Officer e e-mail:',
        'email_compliance': 'E-mail do Compliance Officer:',
        'codigo_conduta_selecionado': 'A empresa possui Código de Conduta/Ética?',
        'pldft_selecionado': 'A Empresa possui políticas e procedimentos de prevenção a fraudes?',
        'organograma_selecionado': 'Existem outras empresas que constam no organograma da organização?',
        'detalhes_organograma': 'Por favor, descreva os detalhes do organograma:',
        'autuacao_selecionado': 'A empresa ou qualquer outra pertencente ao mesmo grupo econômico já foi autuada?',
        'detalhes_autuacao': 'Detalhes da Autuação:',
        'outra_sociedade_selecionado': 'Os sócios fazem parte de outras sociedades/empresas?',
        'detalhes_outra_sociedade': 'Detalhes da Outra Sociedade:',
        'processos_selecionado': 'A empresa ou seus sócios já foram objeto de processos administrativos ou criminais?',
        'detalhes_processos': 'Detalhes dos Processos:',
        'pep_selecionado': 'A empresa ou seus sócios são considerados PEP (Pessoas Expostas Politicamente)?',
        'detalhes_pep': 'Detalhes PEP:',
        'citado_midia_selecionado': 'A empresa e/ou seus sócios, é (são) citada(os) em mídia por suspeitas de atividades criminais?',
        'percentual_dinheiro_selecionado': 'Qual o percentual que representa dinheiro em espécie?',
        'gestao_empresa': 'Por quem é feita a gestão da empresa?',
        'atividades_empresa': 'Qual(is) a(s) atividade(s) da empresa?',
        'origem_recursos': 'Se a empresa foi constituída há menos de 2 anos, informar a origem dos recursos:',
        'fluxo_dinheiro': 'Qual o fluxo do dinheiro? Como vai operar cash-in e cash-out?',
        'volume_operacao': 'Qual o volume esperado da operação?',
        'maturidade_tecnologica': 'Qual a maturidade tecnológica da empresa? (Iniciante, intermediário ou avançado)',
        'uso_3xpay': 'Por que contratou ou pretende contratar o 3X Pay e como vai utilizar nossas soluções?',
        'nome_representante': 'Nome do Representante Legal do Contrato:',
        'cargo_representante': 'Cargo:',
        'telefone_representante': 'Telefone:',
        'email_representante': 'E-mail:',
        'cpf_representante': 'CPF:',
        'rg_representante': 'RG:',
        'servicos': 'Objeto da Parceria: Quais serviços do 3X Pay serão utilizados? (Transações via Pix, Transações Cartão de Crédito, APIs, Emissão de Boletos)'
    }

    for i in range(2, 6):
        question_mapping[f'nome_socio{i}'] = f'Nome Completo dos Sócios/Diretores/Administradores/Representantes Legais {i}:'
        question_mapping[f'cpf_socio{i}'] = f'CPF/CNPJ {i}:'
        question_mapping[f'data_nascimento{i}'] = f'Data de Nascimento {i}:'
        question_mapping[f'percentual_societario{i}'] = f'Percentual Societário {i}:'
        question_mapping[f'us_person{i}'] = f'É US Person {i}?'

    return question_mapping.get(key, key.replace('_', ' ').title())

def format_value(value, key):
    if value.lower() == "sim":
        return "Sim"
    elif value.lower() == "não":
        return "Não"
    return value

def save_to_word(data):
    document = Document()
    
    title = document.add_heading('Relatório de Dados do Formulário', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(14)
    run.font.name = 'Arial Narrow'

    document.add_paragraph()  # Espaçamento

    nome_arquivo = f"Relatório {data.get('nome_socio1', 'Relatório')}.docx"

    for key, value in data.items():
        formatted_key = format_key(key)
        formatted_value = format_value(value, key)

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{formatted_key}")
        run.bold = True
        run.font.name = 'Arial Narrow'
        
        response_paragraph = document.add_paragraph(formatted_value)
        response_paragraph.style.font.name = 'Arial Narrow'

    document.save(nome_arquivo)

# Configuração da interface gráfica
root = tk.Tk()
root.geometry("600x500")  # Tamanho da janela
root.title("Formulário para Word")
root.configure(bg="#f7f7f7")  # Cor de fundo da janela

# Frame principal
frame = tk.Frame(root, bg="#f7f7f7")
frame.pack(pady=20)

# Texto de instrução
instruction_label = tk.Label(frame, text="Insira os dados abaixo:", bg="#f7f7f7", font=("Arial", 12, "bold"))
instruction_label.pack()

# Área de texto
text_area = tk.Text(frame, height=20, width=70, bg="#ffffff", fg="#333333", font=("Arial", 12), borderwidth=2, relief="solid")
text_area.pack(pady=10)

# Botão
process_button = tk.Button(frame, text="Exportar para Word", command=process_text, bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), relief="raised")
process_button.pack(pady=10)

# Executar a aplicação
root.mainloop()